from selenium import webdriver
from selenium.webdriver.common.alert import Alert
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By

from selenium.webdriver.common.keys import Keys

import time
import docx


def checkplag(path):

    submitstatus = driver.find_element(
        By.XPATH, '//*[@id="assignment_132050245"]/td[1]/div')

    submitparent = submitstatus.find_element(By.XPATH, "..")

    submitrowparent = submitparent.find_element(By.XPATH, "..")

    status = submitrowparent.find_element(By.CLASS_NAME, "btn-primary")

    submission = False
    if status.text == "Submit":
        submission = True
        home_page_source = driver.page_source
        soup = BeautifulSoup(home_page_source, 'html.parser')
        links = soup.findAll('a', {'class': 'btn btn-primary submit-btn'})
        submitlink = links[0]["href"]
        submit = driver.find_element(
            By.XPATH, "//a[@href='" + submitlink + "']")

    else:
        home_page_source = driver.page_source
        soup = BeautifulSoup(home_page_source, 'html.parser')
        submitbutton = soup.findAll(
            'a', {'class': 'btn btn-primary tooltip submit-btn'})
        # print(submitbutton)
        submitlink = submitbutton[5]["XPATH"]
        submit = driver.find_element(
            By.XPATH, "//a[@href='" + submitlink + "']")
        
        # print(links)
        # submit = driver.find_element(By.PARTIAL_LINK_TEXT, "Resubmit")

    print("uploading the file....")

    parent = submit.find_element(By.XPATH, "..")

    rowparent = parent.find_element(By.XPATH, "..")

    rowid = rowparent.get_attribute("id")

    if status.text == "Submit":
        driver.get("https://www.turnitin.com/"+submitlink)
    else:
        submit.click()
        driver.get(driver.current_url)

    time.sleep(5)

    if not submission:
        alert = Alert(driver)
        alert.accept()

    time.sleep(10)

    driver.find_element("name", "title").send_keys("assignment tester")
    driver.find_element("id", "selected-file").send_keys(path)

    driver.find_element("id", "upload-btn").click()

    time.sleep(10)

    driver.find_element(By.ID, "confirm-btn").click()

    time.sleep(10)

    driver.find_element("id", "close-btn").click()

    get_url = driver.current_url

    driver.get(get_url)
    print("file uploaded. Calculating plagiarism")
    time.sleep(10)
    plag = driver.find_element(By.ID, rowid).find_element(
        By.CLASS_NAME, "or-percentage").text

    filedownload = driver.find_element(By.ID, rowid).find_element(
        By.CLASS_NAME, "download")
    filedownload.click()
    driver.find_element(By.PARTIAL_LINK_TEXT, "PDF format")
    viewbutton = driver.find_element(
        By.ID, rowid).find_element(By.PARTIAL_LINK_TEXT, "View")
    viewbutton.click()

    print("Plagiarism in the uploaded part is " + plag)
    plag = plag[:-1]
    plag = int(plag)
    return plag


def readtxt(filename):
    doc = docx.Document(filename)
    if len(doc.paragraphs) > 0:
        last_paragraph = doc.paragraphs[-1]
    else:
        print("Document is empty")
    fullText = " "
    i = 0
    for para in doc.paragraphs:
        fullText += para.text
        words = fullText.split()
        if (len(words) > 150):
            i = i+1
            new_doc = docx.Document()
            new_doc.add_paragraph(fullText)
            new_doc.save(f"D://output_{i}.docx")
            docnames.append(f"D://output_{i}.docx")
            fullText = " "


docnames = []
readtxt('original_document.docx')


options = webdriver.ChromeOptions()
# options.headless = True
driver = webdriver.Chrome(chrome_options=options)
driver.get("https://www.turnitin.com/login_page.asp?")

driver.maximize_window()

print("Entering credentials")

driver.find_element("name", "email").send_keys("........")###### enter your own mail

driver.find_element("name", "user_password").send_keys("...........") ###### enter your own pasword

driver.find_element("name", "submit").click()

get_url = driver.current_url
print("logging in")

driver.get(get_url)

driver.find_element(By.XPATH,
                    '//*[@id="main_section"]/div/div[2]/table/tbody/tr[3]/td[2]/a').click()


get_url = driver.current_url


driver.get(get_url)


plagsum = 0

for path in docnames:
    plagFile = checkplag(path)
    plagsum = plagsum+plagFile

files = len(docnames)
print("TOTAL PLAGIARISM IN FILE IS", plagsum // files)
